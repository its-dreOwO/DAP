"""One-shot script to generate project_planning.docx in docs/."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── styles ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        shading.set(qn("w:color"), "FFFFFF")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    return table


# ── Title page ───────────────────────────────────────────────────────────────
title = doc.add_heading("Project Execution Plan", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(
    "DAP391m — Project 8: Supplier Lead-Time Risk Prediction for Retail Procurement"
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

meta = doc.add_paragraph(
    "FPT University HCMC  ·  Summer 2026  ·  Group 8\n"
    "Nguyễn Hoài Khánh  ·  Hồ Lâm Bảo Đăng  ·  Dương Gia Bảo\n"
    "Supervisor: Mr. Nguyen Hoai Linh"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Current State ────────────────────────────────────────────────────────────
add_heading("Current Project State", level=1)
add_para(
    "Before execution begins, the repository is scaffolded as follows:", italic=True
)
add_table(
    ["Artifact", "Status"],
    [
        ["report/main.tex", "Skeleton present — Results & Conclusion are TODO stubs"],
        ["src-code/01_ingestion_cleaning.py", "Partially written (28 lines)"],
        ["src-code/02_sql_analysis.py", "Empty stub"],
        ["src-code/03_eda.py", "Empty stub"],
        ["src-code/04_feature_engineering.py", "Empty stub"],
        ["src-code/05_modeling.py", "Empty stub"],
        ["src-code/06_visualization_advanced.py", "Empty stub"],
        ["app/app.py", "Streamlit UI skeleton exists"],
        ["Data/filtered/model_features.csv", "Generated — ready to use"],
        ["sql/analysis.sql", "Scaffolded"],
        ["docs/AI_AuditLog_Template_DAP391m.xlsx", "Template present — not filled"],
    ],
)

doc.add_paragraph()

# ── Execution Order ───────────────────────────────────────────────────────────
add_heading("Execution Order & Dependencies", level=1)
add_para("Scripts must run in the following sequence. The critical path is bolded.")
add_table(
    ["Step", "Script / Deliverable", "Depends On", "Critical Path"],
    [
        ["1", "01_ingestion_cleaning.py", "Raw CSVs", "YES"],
        ["2", "02_sql_analysis.py", "Step 1", "No"],
        ["3", "03_eda.py", "Step 1", "No"],
        ["4", "04_feature_engineering.py", "Step 1", "YES"],
        ["5", "05_modeling.py", "Step 4", "YES"],
        ["6", "06_visualization_advanced.py", "Steps 4+5", "No"],
        ["7", "app/app.py", "Step 5 (needs .pkl)", "YES"],
        ["8", "Power BI dashboard", "Steps 1, 2, 5", "No"],
        ["9", "report/main.tex (fill Results)", "Step 5 (metrics)", "YES"],
        ["10", "AI Audit Log", "Throughout project", "No"],
    ],
)

doc.add_paragraph()

# ── Phase 1 ───────────────────────────────────────────────────────────────────
add_heading("Phase 1 — Data Pipeline (Scripts 01–04)", level=1)

add_heading("1.1  01_ingestion_cleaning.py  (partial → complete)", level=2)
add_bullet("Load all three raw CSVs with explicit dtype and parse_dates arguments")
add_bullet(
    "Audit nulls per column → impute numeric nulls with median, drop rows with null target"
)
add_bullet("Fix supplier_id type consistency across customer.csv and shipment.csv")
add_bullet(
    "Build country_to_region dict from unique values in shipment.D_Country → map to Asia / Europe / Americas / Middle East / Africa"  # noqa: E501
)
add_bullet(
    "Perform inner joins: shipment ← customer on supplier_id, then ← logistics_performance on region"
)
add_bullet(
    "Assert final shape is exactly 704 rows × 23 columns; raise ValueError if not"
)
add_bullet(
    "Save: Data/filtered/customer_clean.csv, shipment_clean.csv, logistics_performance_clean.csv"
)

doc.add_paragraph()
add_heading("1.2  02_sql_analysis.py  (empty → full build)", level=2)
add_bullet("Load model_features.csv into an in-memory SQLite database via SQLAlchemy")
add_bullet(
    "Execute and print all 6 required queries; save each result as a CSV in Data/filtered/sql_outputs/"
)
add_bullet("Query 1: Average lead time per supplier — GROUP BY supplier_id")
add_bullet("Query 2: Delay frequency per supplier — SUM(is_delayed) / COUNT(*)")
add_bullet(
    "Query 3: Volume–delay correlation — Pearson on grouped aggregates via pandas"
)
add_bullet("Query 4: Monthly delay trends — strftime group by YYYY-MM")
add_bullet("Query 5: Year-over-year growth rate — self-join or LAG window on year")
add_bullet(
    "Query 6: Carrier penetration index — (delay_share / volume_share) per carrier"
)
add_bullet(
    "Mirror all queries to sql/analysis.sql as static reference (course deliverable)"
)

doc.add_paragraph()
add_heading("1.3  03_eda.py  (empty → full build)", level=2)
add_bullet("Class balance bar chart (16% Delayed vs 84% On-Time) → PNG")
add_bullet("Correlation heatmap of all 23 numeric features → PNG")
add_bullet(
    "Distribution plots for 5 key numeric features (lead_time_days, customs_clearance_time_days, freight_cost, delay_hours_avg, warehouse_utilization_percent)"  # noqa: E501
)
add_bullet("Time-series line chart: monthly delayed shipment count")
add_bullet("Boxplot: lead_time_days split by delivery_status")
add_bullet("All figures saved to docs/figures/eda/")

doc.add_paragraph()
add_heading("1.4  04_feature_engineering.py  (empty → full build)", level=2)
add_bullet(
    "Label-encode categoricals: carrier, market_segment, region — save encoder mappings as JSON to Data/filtered/encoders.json (Streamlit reuse)"  # noqa: E501
)
add_bullet(
    "StandardScaler on all numeric features — save fitted scaler to Data/filtered/scaler.pkl"
)
add_bullet("Stratified 80/20 train/test split on delivery_status target")
add_bullet(
    "Save X_train.csv, X_test.csv, y_train.csv, y_test.csv to Data/filtered/splits/"
)
add_bullet(
    "Print class distribution in both splits to confirm stratification is correct"
)

doc.add_paragraph()

# ── Phase 2 ───────────────────────────────────────────────────────────────────
add_heading("Phase 2 — Modeling (Script 05)", level=1)

add_heading("2.1  05_modeling.py  (empty → full build — most critical script)", level=2)

add_para("Model training:", bold=True)
add_bullet("Load train/test splits from Data/filtered/splits/")
add_bullet('LogisticRegression(class_weight="balanced", max_iter=1000)')
add_bullet('DecisionTreeClassifier(class_weight="balanced")')
add_bullet('RandomForestClassifier(class_weight="balanced", n_estimators=200)')
add_bullet(
    'XGBClassifier(scale_pos_weight≈5.25, eval_metric="aucpr")  ← scale_pos_weight = count(On-Time)/count(Delayed)'
)

add_para("Evaluation per model:", bold=True)
add_bullet("Stratified 5-fold CV → mean ± std for PR-AUC and Recall(Delayed)")
add_bullet("Held-out test set metrics: PR-AUC, Recall, ROC-AUC, F1, confusion matrix")
add_bullet(
    "Threshold tuning: scan 0.1–0.9, select threshold maximising F1 on Delayed class"
)
add_bullet("Print consolidated table matching LaTeX tab:results format exactly")

add_para("SHAP (XGBoost only):", bold=True)
add_bullet(
    "TreeExplainer → global summary_plot (beeswarm) → save docs/figures/shap_summary.png"
)
add_bullet(
    "Waterfall plot for highest-risk shipment in test set → save docs/figures/shap_waterfall_example.png"
)

add_para("Artifacts saved:", bold=True)
add_bullet("Data/filtered/xgb_model.pkl  (consumed by Streamlit app)")
add_bullet("Data/filtered/optimal_threshold.txt")
add_bullet("Data/filtered/model_metrics.json  (all 4 models — report reads this)")

doc.add_paragraph()

# ── Phase 3 ───────────────────────────────────────────────────────────────────
add_heading("Phase 3 — Visualisation (Script 06)", level=1)

add_heading("3.1  06_visualization_advanced.py  (empty → full build)", level=2)
add_bullet(
    "Folium HeatMap: lat/lon of destination countries, intensity = XGBoost delay probability → export docs/figures/risk_heatmap.html"  # noqa: E501
)
add_bullet(
    "Plotly boxplots: lead_time_days grouped by top-20 suppliers by volume → export docs/figures/supplier_boxplot.html"
)
add_bullet("Lead-time histogram: overlay Delayed vs On-Time distributions → export PNG")
add_bullet(
    "Monthly trend line: Plotly dual-axis (shipment count + delay rate) → export HTML"
)
add_bullet("All interactive charts also exported as static PNGs for the LaTeX report")

doc.add_paragraph()

# ── Phase 4 ───────────────────────────────────────────────────────────────────
add_heading("Phase 4 — Streamlit App (app/app.py)", level=1)

add_heading("4.1  Complete the existing skeleton", level=2)
add_bullet(
    "Load xgb_model.pkl and scaler.pkl at startup with @st.cache_resource to avoid reload on every interaction"
)
add_bullet(
    "Add carrier and region dropdowns to sidebar — apply saved encoders.json to convert to integer codes"
)
add_bullet(
    "On Predict: scale input row using loaded scaler → model.predict_proba → read optimal_threshold.txt → classify as Green/Amber/Red"  # noqa: E501
)
add_bullet("Display gauge or colour-coded st.metric card for delay probability")
add_bullet(
    "SHAP waterfall: TreeExplainer on single input row → st.pyplot waterfall chart"
)
add_bullet(
    'Second tab "Supplier Risk Ranking": load supplier_risk_ranking_final_latest.csv → sortable st.dataframe + Plotly bar chart'  # noqa: E501
)
add_bullet(
    "Test locally: streamlit run app/app.py — verify all 6 sidebar inputs, prediction, and SHAP chart render without errors"  # noqa: E501
)

doc.add_paragraph()

# ── Phase 5 ───────────────────────────────────────────────────────────────────
add_heading("Phase 5 — Power BI Dashboard", level=1)

add_para(
    "Input files: model_features.csv, supplier_risk_ranking_final_latest.csv, all SQL output CSVs from Phase 1.2",
    italic=True,
)

add_heading("5.1  Dashboard pages", level=2)
add_table(
    ["Page", "Visuals", "Key Slicers"],
    [
        [
            "Supplier Scorecard",
            "KPI cards: total shipments, % delayed, avg lead time; conditional formatting red if delay_rate > 20%",
            "supplier_id",
        ],
        [
            "Regional Risk Map",
            "Filled map coloured by delay rate per region",
            "region, date range",
        ],
        [
            "Carrier Performance",
            "Clustered bar: On-Time vs Delayed per carrier; penetration index line overlay",
            "carrier",
        ],
        ["Trend Analysis", "Monthly delay rate trend line; YoY growth table", "Year"],
        [
            "Risk Alerts",
            "Top-10 highest-risk suppliers table with amber/red conditional flags",
            "threshold slider",
        ],
    ],
)
add_para("Save as: docs/powerbi/supplier_scorecard.pbix", italic=True)

doc.add_paragraph()

# ── Phase 6 ───────────────────────────────────────────────────────────────────
add_heading("Phase 6 — Final Report (report/main.tex)", level=1)

add_heading("6.1  Sections requiring content after experiments", level=2)
add_table(
    ["Section", "What to write"],
    [
        [
            "§5 Results",
            "Copy actual metric values from model_metrics.json into tab:results; add 1–2 sentences on why XGBoost wins (PR-AUC delta)",  # noqa: E501
        ],
        [
            "§6 Visualisation",
            "Reference each figure file path with \\includegraphics; 1 sentence per chart describing the procurement insight",  # noqa: E501
        ],
        [
            "§8 Conclusion",
            "3 bullet answers to the 3 research questions; limitations (dataset size, synthetic labels); future work (real-time feed, more carriers, online learning)",  # noqa: E501
        ],
        [
            "References",
            "Add BibTeX entries: scikit-learn, XGBoost paper, SHAP paper, Kaggle dataset DOI, imbalanced-learn",
        ],
    ],
)
add_para(
    "Compile: run pdflatex report/main.tex twice (second pass resolves cross-references).",
    italic=True,
)

doc.add_paragraph()

# ── Phase 7 ───────────────────────────────────────────────────────────────────
add_heading("Phase 7 — AI Audit Log (docs/AI_AuditLog_Template_DAP391m.xlsx)", level=1)

add_heading("7.1  Requirements", level=2)
add_bullet(
    "Document 15–20 prompts used with Claude / ChatGPT during the project (one row per prompt)"
)
add_bullet(
    "Mark at least 3 hallucination checks — cases where AI output was verified and corrected"
)
add_table(
    ["Column", "Description"],
    [
        ["Prompt", "Exact or paraphrased prompt sent to AI"],
        ["Tool Used", "Claude Code / ChatGPT / Copilot etc."],
        ["Output Summary", "What the AI returned"],
        ["Was Output Verified?", "Yes / No"],
        ["Correction Made", "What was wrong and how it was fixed"],
    ],
)

doc.add_paragraph()

# ── Risks ──────────────────────────────────────────────────────────────────────
add_heading("Open Risks & Mitigations", level=1)
add_table(
    ["Risk", "Impact", "Mitigation"],
    [
        [
            "country_to_region dict not yet built",
            "High — joins fail without it",
            "Inspect unique values in shipment.D_Country before writing 01; build exhaustive mapping",
        ],
        [
            "Scaler/encoder mismatch between 04 and app",
            "High — predictions will be wrong",
            "App must load the exact scaler.pkl and encoders.json saved by 04, never retrain inline",
        ],
        [
            "XGBoost scale_pos_weight miscalculated",
            "Medium — biased threshold",
            "Compute from y_train: (y==0).sum() / (y==1).sum() ≈ 5.25",
        ],
        [
            "SHAP version API break",
            "Medium — waterfall chart errors",
            "Pin shap==0.44 in .venv; test shap.plots.waterfall before Streamlit integration",
        ],
        [
            "Report TODO stubs left unfilled",
            "High — submission incomplete",
            "Fill Results section immediately after 05_modeling.py runs; Conclusion last",
        ],
        [
            "Power BI not available on Linux",
            "Medium — dashboard deliverable blocked",
            "Use Power BI Desktop on Windows VM or export charts to PDF/PNG as fallback",
        ],
    ],
)

doc.add_paragraph()

# ── Weekly timeline ────────────────────────────────────────────────────────────
add_heading("Weekly Timeline (Course Schedule)", level=1)
add_table(
    ["Week", "Focus", "Phases"],
    [
        [
            "Week 1",
            "Data understanding, cleaning, SQL analysis, EDA",
            "Phase 1 (01–03)",
        ],
        [
            "Week 2",
            "Feature engineering, all 4 models, SHAP, evaluation",
            "Phases 1.4 + 2",
        ],
        [
            "Week 3",
            "Plotly/Folium visualisation, odds-ratio regression analysis",
            "Phase 3",
        ],
        ["Week 4", "Power BI supplier scorecard", "Phase 5"],
        [
            "Weeks 4–5",
            "Streamlit web application, report write-up, AI Audit Log",
            "Phases 4 + 6 + 7",
        ],
    ],
)

doc.add_paragraph()
add_para(
    "Generated by Claude Code — DAP391m Group 8, FPT University HCMC, Summer 2026",
    italic=True,
)

# ── Save ───────────────────────────────────────────────────────────────────────
out = "/home/dre/Desktop/study/DAP/docs/project_planning.docx"
doc.save(out)
print(f"Saved: {out}")
